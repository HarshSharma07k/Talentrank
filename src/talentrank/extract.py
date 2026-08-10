"""Resume file text extraction (PDF / DOCX) for TalentRank. See enhancements/06.

Deliberately a separate concern from `/match`, not a preprocessing step folded into
it: the user must see and edit the extracted text before matching (pypdf/python-docx
extraction is lossy, especially for multi-column resumes), it keeps `/match`'s cache
key a hash of *text* rather than of file bytes, and extraction failures get their own
clear error messages instead of being buried inside a match failure.
"""

from __future__ import annotations

from dataclasses import dataclass
import io

from docx import Document
import pypdf
from pypdf.errors import PyPdfError

from src.talentrank.config import get_settings

SUPPORTED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"


class UnsupportedFileTypeError(Exception):
    """Neither the filename extension nor the magic bytes matched a supported type,
    or the two disagreed. Maps to 415."""


class EncryptedPdfError(Exception):
    """A password-protected PDF. Maps to 422."""


class EmptyExtractionError(Exception):
    """Parsed successfully (or failed to parse) but yielded no usable text -- a
    scanned/image-only PDF, or a genuinely malformed file. Maps to 422."""


@dataclass(slots=True)
class ExtractionResult:
    text: str
    char_count: int
    page_count: int | None  # None for DOCX
    truncated: bool


def _sniff_kind(filename: str, content: bytes) -> str:
    """Dispatch on extension **and** magic bytes -- never trust the client-supplied
    content type alone (see enhancements/06's Do NOT). A `.pdf`-named file that isn't
    actually a PDF (or vice versa) is rejected rather than handed to the wrong parser.
    """

    lower_name = filename.lower()
    is_pdf_magic = content.startswith(_PDF_MAGIC)
    is_docx_magic = content.startswith(_DOCX_MAGIC)

    if lower_name.endswith(".pdf") and is_pdf_magic:
        return "pdf"
    if lower_name.endswith(".docx") and is_docx_magic:
        return "docx"

    raise UnsupportedFileTypeError(
        f"Unsupported or mismatched file type for {filename!r}. Upload a PDF or DOCX resume."
    )


def _extract_pdf(content: bytes) -> tuple[str, int]:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            raise EncryptedPdfError("This PDF is password-protected. Remove the password or paste the text instead.")
        pages_text = [page.extract_text() or "" for page in reader.pages]
        page_count = len(reader.pages)
    except EncryptedPdfError:
        raise
    except PyPdfError as exc:
        raise EmptyExtractionError(f"Could not read this PDF ({exc}); try pasting the text instead.") from exc

    return "\n".join(pages_text), page_count


def _extract_docx(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]

    # Skills are very often laid out in a table, not prose -- paragraph text alone
    # would silently drop them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def _normalize_extracted_text(text: str) -> str:
    """The same whitespace collapsing `/match` applies before embedding, so what the
    user sees in the textarea is exactly what gets embedded."""

    return " ".join(text.split())


def extract_text_from_upload(filename: str, content: bytes) -> ExtractionResult:
    """Extract, normalize, and length-guard resume text from an uploaded PDF or DOCX.

    Raises `UnsupportedFileTypeError`, `EncryptedPdfError`, or `EmptyExtractionError`
    for the caller (`api.py`) to map to 415/422/422 respectively. Does not enforce
    `max_upload_bytes` -- that is a request-size concern the caller checks before
    ever reading the full body into `content`.
    """

    settings = get_settings()
    kind = _sniff_kind(filename, content)

    if kind == "pdf":
        raw_text, page_count = _extract_pdf(content)
    else:
        raw_text, page_count = _extract_docx(content), None

    normalized = _normalize_extracted_text(raw_text)

    if len(normalized) < settings.min_resume_chars:
        raise EmptyExtractionError(
            "This looks like a scanned image or has no extractable text; try pasting the text instead."
        )

    truncated = len(normalized) > settings.max_resume_chars
    final_text = normalized[: settings.max_resume_chars]

    return ExtractionResult(
        text=final_text,
        char_count=len(final_text),
        page_count=page_count,
        truncated=truncated,
    )
