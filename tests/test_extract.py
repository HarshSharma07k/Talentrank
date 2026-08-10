"""Tests for `extract.py` and `POST /extract-text`. See enhancements/06.

`pypdf`/`python-docx` can *read* arbitrary compliant files but neither has a
convenient "draw some text" API, so the PDF fixtures below hand-build a minimal
single-page PDF (a well-known, short, standards-compliant pattern) rather than
pulling in a new dependency like reportlab just for tests.
"""

from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient
import pypdf
import pytest

from src.talentrank.config import get_settings
from src.talentrank.extract import (
    EmptyExtractionError,
    EncryptedPdfError,
    UnsupportedFileTypeError,
    extract_text_from_upload,
)

_PDF_CONTENT_TYPE = "application/pdf"
_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_minimal_pdf(text: str) -> bytes:
    """A minimal, hand-built single-page PDF with one text-drawing operator.
    `pypdf` tolerates a plain, non-cross-reference-stream xref table like this one."""

    stream = f"BT /F1 24 Tf 20 100 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 200 200] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode())
        out.write(obj)
        out.write(b"\nendobj\n")

    xref_offset = out.tell()
    count = len(objects) + 1
    out.write(f"xref\n0 {count}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        out.write(f"{offset:010d} 00000 n \n".encode())
    out.write(b"trailer\n")
    out.write(f"<< /Size {count} /Root 1 0 R >>\n".encode())
    out.write(b"startxref\n")
    out.write(str(xref_offset).encode())
    out.write(b"\n%%EOF")
    return out.getvalue()


def _build_blank_pdf() -> bytes:
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _build_encrypted_pdf() -> bytes:
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt(user_password="secret", owner_password="secret")
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _build_docx(paragraphs: list[str], table_cells: list[str] | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_cells:
        table = document.add_table(rows=1, cols=len(table_cells))
        for cell, text in zip(table.rows[0].cells, table_cells):
            cell.text = text
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --- extract_text_from_upload (function-level) --------------------------------


def test_pdf_roundtrip() -> None:
    pdf_bytes = _build_minimal_pdf("Jane Doe Software Engineer with Five Years Experience")

    result = extract_text_from_upload("resume.pdf", pdf_bytes)

    assert "Jane Doe Software Engineer with Five Years Experience" in result.text
    assert result.page_count == 1
    assert result.char_count == len(result.text)
    assert result.truncated is False


def test_docx_roundtrip() -> None:
    docx_bytes = _build_docx(["Jane Doe, Software Engineer with five years experience"])

    result = extract_text_from_upload("resume.docx", docx_bytes)

    assert "Jane Doe, Software Engineer with five years experience" in result.text
    assert result.page_count is None
    assert result.truncated is False


def test_docx_table_text_included() -> None:
    """Skills are very often laid out in a table, not prose."""

    docx_bytes = _build_docx(["Summary paragraph."], table_cells=["Python", "FastAPI", "Docker"])

    result = extract_text_from_upload("resume.docx", docx_bytes)

    assert "Python" in result.text
    assert "FastAPI" in result.text
    assert "Docker" in result.text


def test_truncation_flag() -> None:
    settings = get_settings()
    long_word = "x" * (settings.max_resume_chars + 500)
    docx_bytes = _build_docx([long_word])

    result = extract_text_from_upload("resume.docx", docx_bytes)

    assert result.truncated is True
    assert result.char_count == settings.max_resume_chars


def test_encrypted_pdf_raises() -> None:
    with pytest.raises(EncryptedPdfError):
        extract_text_from_upload("resume.pdf", _build_encrypted_pdf())


def test_empty_extraction_raises() -> None:
    with pytest.raises(EmptyExtractionError):
        extract_text_from_upload("resume.pdf", _build_blank_pdf())


def test_unsupported_type_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload("resume.txt", b"just plain text, not a real resume file")


def test_magic_bytes_mismatch_raises() -> None:
    """A file named `.pdf` whose bytes are not actually a PDF."""

    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload("resume.pdf", b"this is not actually a pdf")


# --- POST /extract-text (API-level, exact status codes) ------------------------


def test_pdf_roundtrip_via_api(client: TestClient) -> None:
    pdf_bytes = _build_minimal_pdf("Jane Doe Software Engineer with Five Years Experience")

    response = client.post("/extract-text", files={"file": ("resume.pdf", pdf_bytes, _PDF_CONTENT_TYPE)})

    assert response.status_code == 200
    body = response.json()
    assert "Jane Doe Software Engineer with Five Years Experience" in body["text"]
    assert body["filename"] == "resume.pdf"
    assert body["page_count"] == 1
    assert body["truncated"] is False


def test_oversize_413(client: TestClient) -> None:
    get_settings().max_upload_bytes = 100

    response = client.post(
        "/extract-text",
        files={"file": ("resume.pdf", _build_minimal_pdf("x" * 500), _PDF_CONTENT_TYPE)},
    )

    assert response.status_code == 413
    assert "detail" in response.json()


def test_wrong_type_415(client: TestClient) -> None:
    txt_response = client.post(
        "/extract-text", files={"file": ("resume.txt", b"plain text resume content", "text/plain")}
    )
    assert txt_response.status_code == 415

    exe_response = client.post(
        "/extract-text", files={"file": ("resume.exe", b"MZ\x90\x00fake-exe-bytes", "application/octet-stream")}
    )
    assert exe_response.status_code == 415


def test_magic_bytes_mismatch_415(client: TestClient) -> None:
    response = client.post(
        "/extract-text",
        files={"file": ("resume.pdf", b"this is not actually a pdf despite the extension", _PDF_CONTENT_TYPE)},
    )

    assert response.status_code == 415


def test_empty_extraction_422(client: TestClient) -> None:
    response = client.post("/extract-text", files={"file": ("resume.pdf", _build_blank_pdf(), _PDF_CONTENT_TYPE)})

    assert response.status_code == 422
    assert "detail" in response.json()


def test_encrypted_pdf_422(client: TestClient) -> None:
    response = client.post("/extract-text", files={"file": ("resume.pdf", _build_encrypted_pdf(), _PDF_CONTENT_TYPE)})

    assert response.status_code == 422
    assert "detail" in response.json()


def test_docx_roundtrip_via_api(client: TestClient) -> None:
    docx_bytes = _build_docx(["Jane Doe, Software Engineer with five years experience"], table_cells=["Python", "SQL"])

    response = client.post("/extract-text", files={"file": ("resume.docx", docx_bytes, _DOCX_CONTENT_TYPE)})

    assert response.status_code == 200
    body = response.json()
    assert "Jane Doe, Software Engineer with five years experience" in body["text"]
    assert "Python" in body["text"]
    assert body["page_count"] is None
